"""Generate HTML and Word reports from grading results."""
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DIMENSION_COLORS = {
    "typos":      {"hex": "#FF9800", "name": "错别字"},
    "grammar":    {"hex": "#E91E63", "name": "语法"},
    "structure":  {"hex": "#4CAF50", "name": "文章结构"},
    "logic":      {"hex": "#2196F3", "name": "逻辑性"},
    "technique":  {"hex": "#9C27B0", "name": "文笔手法"},
    "vocabulary": {"hex": "#FDD835", "name": "词汇运用"},
}

DIMENSION_LABELS_ZH = {k: v["name"] for k, v in DIMENSION_COLORS.items()}


def export_html(original_text: str, result: dict) -> str:
    """Generate a standalone HTML report with dual-pane comment layout."""
    annotations = result.get("annotations", [])
    essay_html = _build_essay_html(original_text, annotations)
    comments_html = _build_comments_html(annotations)

    return f"""<!DOCTYPE html>
<html lang="{result.get('language', 'zh')}">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>作文批改报告</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif; background: #f0f2f5; color: #333; }}
.container {{ max-width: 1100px; margin: 0 auto; padding: 24px; }}
.header {{ background: #fff; border-radius: 12px; padding: 24px; margin-bottom: 16px; box-shadow: 0 1px 3px rgba(0,0,0,0.08); }}
.header h1 {{ font-size: 1.4rem; margin-bottom: 8px; }}
.meta {{ color: #888; font-size: 0.85rem; }}
.overall {{ background: #f0f7ff; border-left: 4px solid #2563eb; border-radius: 6px; padding: 12px 16px; margin-bottom: 16px; font-size: 0.92rem; }}
.main {{ display: flex; gap: 0; background: #fff; border-radius: 12px; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,0.08); }}
.essay-panel {{ flex: 1; padding: 28px 24px; line-height: 2.2; font-size: 0.95rem; border-right: 1px solid #eee; overflow-y: auto; max-height: 70vh; }}
.essay-panel p {{ margin-bottom: 14px; text-indent: 2em; }}
.essay-panel .hl {{ border-radius: 3px; padding: 1px 2px; cursor: pointer; transition: background 0.2s; }}
.essay-panel .hl-num {{ display: inline-block; color: #fff; font-size: 0.6rem; border-radius: 50%; width: 16px; height: 16px; text-align: center; line-height: 16px; margin-left: 1px; vertical-align: super; font-weight: 600; }}
.comments-panel {{ width: 340px; background: #fafafa; padding: 20px 16px; overflow-y: auto; max-height: 70vh; }}
.comments-panel h3 {{ font-size: 0.9rem; color: #888; margin-bottom: 14px; }}
.comment-card {{ background: #fff; border-radius: 8px; padding: 12px 14px; margin-bottom: 10px; box-shadow: 0 1px 2px rgba(0,0,0,0.04); transition: box-shadow 0.2s; }}
.comment-card:hover {{ box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
.comment-card.target {{ box-shadow: 0 0 0 2px #2563eb; }}
.comment-header {{ display: flex; align-items: center; gap: 8px; margin-bottom: 6px; }}
.comment-num {{ display: inline-flex; align-items: center; justify-content: center; width: 20px; height: 20px; border-radius: 50%; color: #fff; font-size: 0.7rem; font-weight: 700; flex-shrink: 0; }}
.comment-dim {{ font-weight: 600; font-size: 0.82rem; }}
.comment-issue {{ color: #666; font-size: 0.8rem; margin-bottom: 4px; }}
.comment-sug {{ color: #333; font-size: 0.84rem; line-height: 1.6; }}
.comment-sug em {{ background: #fff9c4; font-style: normal; padding: 1px 3px; border-radius: 2px; }}
@media (max-width: 768px) {{ .main {{ flex-direction: column; }} .comments-panel {{ width: 100%; }} }}
@media print {{ body {{ background: #fff; }} .container {{ max-width: 100%; padding: 0; }} }}
</style>
</head>
<body>
<div class="container">
<div class="header">
  <h1>📝 作文批改报告</h1>
  <div class="meta">字数：{result.get('word_count', '—')} · 语言：{'中文' if result.get('language') == 'zh' else 'English'}</div>
</div>
<div class="overall">
  <strong>📊 总评：</strong>{result.get('overall_comment', '')}
</div>
<div class="main">
  <div class="essay-panel">{essay_html}</div>
  <div class="comments-panel">
    <h3>💬 批注建议（{len(annotations)} 条）</h3>
    {comments_html}
  </div>
</div>
</div>
</body>
</html>"""


def _build_essay_html(text: str, annotations: list) -> str:
    """Insert highlighted spans and annotation number markers into essay text."""
    indexed = []
    for ann in annotations:
        pos = text.find(ann["highlight_text"])
        indexed.append((pos if pos >= 0 else 99999, ann))
    indexed.sort(key=lambda x: x[0])

    result = text
    for i, (_, ann) in enumerate(indexed):
        color = DIMENSION_COLORS.get(ann["dimension"], {}).get("hex", "#999")
        hl = ann["highlight_text"]
        if hl in result:
            result = result.replace(
                hl,
                f'<span class="hl" style="background:{color}33;border-bottom:2px solid {color};" title="{ann["dimension"]} #{ann["id"]}">{hl}</span><span class="hl-num" style="background:{color};">{ann["id"]}</span>',
                1
            )

    paragraphs = result.split('\n')
    para_html = []
    for p in paragraphs:
        p = p.strip()
        if p:
            para_html.append(f'<p>{p}</p>')
    return '\n'.join(para_html)


def _build_comments_html(annotations: list) -> str:
    """Build HTML for the comment cards in the right panel."""
    cards = []
    for ann in annotations:
        dim = ann.get("dimension", "")
        info = DIMENSION_COLORS.get(dim, {"hex": "#999", "name": dim})
        suggestion = ann.get("suggestion", "").replace("\n", "<br>")
        cards.append(f"""<div class="comment-card" id="comment-{ann['id']}">
  <div class="comment-header">
    <span class="comment-num" style="background:{info['hex']};">{ann['id']}</span>
    <span class="comment-dim" style="color:{info['hex']};">{info['name']}</span>
  </div>
  <div class="comment-issue">📍 {ann.get('location', '')} — {ann.get('issue', '')}</div>
  <div class="comment-sug">💡 {suggestion}</div>
</div>""")
    return '\n'.join(cards)


def export_docx(original_text: str, result: dict) -> bytes:
    """Generate a Word document with native Comments for each annotation."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    title = doc.add_heading('作文批改报告', level=1)

    lang = '中文' if result.get('language') == 'zh' else 'English'
    doc.add_paragraph(f"字数：{result.get('word_count', '—')} · 语言：{lang}")

    doc.add_heading('总评', level=2)
    doc.add_paragraph(result.get('overall_comment', ''))

    doc.add_heading('作文原文', level=2)
    doc.add_paragraph(original_text)

    doc.add_heading(f'批注建议（{len(result.get("annotations", []))} 条）', level=2)

    for ann in result.get("annotations", []):
        dim = DIMENSION_COLORS.get(ann.get("dimension", ""), {})
        dim_name = dim.get("name", ann.get("dimension", ""))

        p = doc.add_paragraph()
        runner = p.add_run(f"#{ann['id']} [{dim_name}] {ann.get('location', '')}")
        runner.bold = True
        runner.font.size = Pt(10)

        doc.add_paragraph(f"问题：{ann.get('issue', '')}", style='List Bullet')
        doc.add_paragraph(f"建议：{ann.get('suggestion', '')}", style='List Bullet')
        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
